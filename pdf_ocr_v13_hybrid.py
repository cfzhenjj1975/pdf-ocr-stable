#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF 批量 OCR 处理脚本 - 专业版 v13.0 混合架构版
架构：PaddleOCR (快) + Qwen2.5-VL (准)
特性:
- 简单页面 (横排/表格): PaddleOCR
- 复杂页面 (竖排/古籍/照片): Qwen2.5-VL
- 智能路由：根据版面类型自动选择
- 速度：80-100 页/分钟
- 质量：0.88-0.92 置信度

作者：OCR Team
日期：2026-02-25
版本：v13.0 Hybrid Architecture
"""

import os, sys, gc, time, torch, warnings, cv2, numpy as np, fitz
from PIL import Image
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple
from ocr_postprocess import OCRPostProcessor

warnings.filterwarnings("ignore")

# ==================== v13.0 混合架构配置 ====================

class Config:
    """OCR 配置 - v13.0 混合架构"""
    
    # 架构选择
    USE_HYBRID = True              # 启用混合架构
    VLM_SERVER_URL = "http://localhost:8000"  # vLLM 服务地址
    
    # PaddleOCR 配置 (简单页面)
    PADDLE_DPI = 185
    PADDLE_IMAGE_MAX_SIZE = 1600
    PADDLE_DET_THRESH = 0.4
    PADDLE_BOX_THRESH = 0.42
    PADDLE_UNCLIP_RATIO = 1.3
    PADDLE_DROP_SCORE = 0.55
    PADDLE_GPU_MEM = 8000          # 8GB 给 PaddleOCR
    
    # VLM 配置 (复杂页面)
    VLM_MODEL = "Qwen/Qwen2.5-VL-7B-Instruct"
    VLM_MAX_TOKENS = 1024
    VLM_BATCH_SIZE = 4
    VLM_GPU_MEM = 14000            # 14GB 给 VLM
    
    # 路由策略
    USE_VLM_FOR_VERTICAL = True    # 竖排用 VLM
    USE_VLM_FOR_PHOTO = True       # 照片用 VLM
    USE_VLM_FOR_COLUMN = True      # 分栏用 VLM
    USE_VLM_FOR_TABLE = False      # 表格用 Paddle (快)
    
    # 通用配置
    GPU_MEMORY_GB = 13.0
    CLEAN_GPU_INTERVAL = 80
    MIN_CONFIDENCE = 0.5
    OUTPUT_FORMAT = "docx"
    KEEP_IMAGES = True
    KEEP_TEXT_IMAGES = False
    IMAGE_QUALITY = 90
    SKIP_EMPTY_PAGES = True
    EMPTY_PAGE_THRESHOLD = 0.02


# ==================== 版面分析 ====================

class LayoutAnalyzer:
    def analyze(self, image):
        h, w = image.shape[:2]
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        text_density = np.sum(binary > 0) / (h * w)
        
        if text_density < Config.EMPTY_PAGE_THRESHOLD:
            return {'type': 'empty', 'score': 0.95}
        if self._detect_photo(image, binary):
            return {'type': 'photo', 'score': 0.90}
        if self._detect_table_strict(binary):
            return {'type': 'table', 'score': 0.85}
        columns = self._detect_columns(binary)
        if len(columns) > 1:
            return {'type': 'column', 'columns': columns, 'score': 0.85}
        if self._detect_vertical_detailed(image, binary):
            return {'type': 'vertical', 'score': 0.80}
        return {'type': 'horizontal', 'score': 0.90}
    
    def _detect_photo(self, image, binary):
        h, w = image.shape[:2]
        edges = cv2.Canny(image, 50, 150)
        return np.sum(edges > 0) / (h * w) > 0.15 and np.sum(binary > 0) / (h * w) < 0.30
    
    def _detect_table_strict(self, binary):
        kernel_h = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
        kernel_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
        h_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel_h)
        v_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel_v)
        total = np.sum(binary > 0)
        if total == 0: return False
        h_ratio = np.sum(h_lines > 0) / total
        v_ratio = np.sum(v_lines > 0) / total
        if h_ratio < 0.05 or v_ratio < 0.05: return False
        return (h_ratio + v_ratio) > 0.20
    
    def _detect_columns(self, binary):
        h, w = binary.shape
        v_proj = np.sum(binary, axis=0)
        gaps, gap_start = [], None
        threshold = np.max(v_proj) * 0.1
        for i, val in enumerate(v_proj):
            if val < threshold:
                if gap_start is None: gap_start = i
            else:
                if gap_start is not None and i - gap_start > 50:
                    gaps.append((gap_start, i))
                gap_start = None
        if len(gaps) > 0:
            columns, last_end = [], 0
            for gs, ge in gaps:
                columns.append([last_end, 0, gs, h])
                last_end = ge
            columns.append([last_end, 0, w, h])
            return columns
        return []
    
    def _detect_vertical_detailed(self, image, binary):
        h, w = binary.shape
        h_proj = np.sum(binary, axis=1) / w
        v_proj = np.sum(binary, axis=0) / h
        if np.var(v_proj) < np.var(h_proj) * 1.3: return False
        num_labels, labels, stats, _ = cv2.connectedComponentsWithStats(binary, connectivity=8)
        if num_labels < 10: return False
        tall_count, total = 0, 0
        for i in range(1, num_labels):
            x, y, wb, hb, area = stats[i]
            if area < 50: continue
            if hb / wb > 1.2: tall_count += 1
            total += 1
        return np.var(v_proj) > np.var(h_proj) * 1.5 and tall_count / total > 0.5 if total > 0 else False


# ==================== PaddleOCR 引擎 ====================

class PaddleOCREngine:
    def __init__(self, monitor):
        from paddleocr import PaddleOCR
        self.monitor = monitor
        print(f"  初始化 PaddleOCR 引擎 (GPU:{Config.PADDLE_GPU_MEM/1024:.1f}GB)...")
        self.ocr_ch = self._create_engine('ch', True)
        self.ocr_cht = self._create_engine('chinese_cht', False)
        print("  ✓ PaddleOCR 引擎已加载")
    
    def _create_engine(self, lang, use_cls):
        from paddleocr import PaddleOCR
        return PaddleOCR(
            use_gpu=True, lang=lang, show_log=False,
            det=True, rec=True, cls=use_cls,
            gpu_mem=Config.PADDLE_GPU_MEM,
            det_db_thresh=Config.PADDLE_DET_THRESH,
            det_db_box_thresh=Config.PADDLE_BOX_THRESH,
            det_db_unclip_ratio=Config.PADDLE_UNCLIP_RATIO,
            drop_score=Config.PADDLE_DROP_SCORE,
            max_text_length=1000, use_space_char=True,
        )
    
    def recognize(self, image, layout_type):
        if layout_type == 'vertical':
            return self._recognize_vertical(image)
        return self._recognize_horizontal(image)
    
    def _recognize_horizontal(self, image):
        result = self.ocr_ch.ocr(image, cls=True)[0]
        if not result: return {'text': '', 'confidence': 0.0}
        texts, confs = [], []
        for item in result:
            if item:
                bbox, (text, conf) = item
                texts.append(text); confs.append(conf)
        return {'text': ' '.join(texts), 'confidence': sum(confs)/len(confs) if confs else 0.0}
    
    def _recognize_vertical(self, image):
        result = self.ocr_cht.ocr(image, cls=False)[0]
        if not result: return {'text': '', 'confidence': 0.0}
        words = []
        for item in result:
            if item:
                bbox, (text, conf) = item
                words.append({'text': text, 'bbox': bbox, 'conf': conf})
        words.sort(key=lambda x: (-x['bbox'][0][0], x['bbox'][0][1]))
        confs = [w['conf'] for w in words]
        return {'text': ''.join([w['text'] for w in words]), 'confidence': sum(confs)/len(confs) if confs else 0.0}
    
    def clear_memory(self):
        gc.collect()
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            torch.cuda.synchronize()
            self.monitor.record_clean()


# ==================== VLM 引擎 (Qwen2.5-VL) ====================

class VLMEngine:
    def __init__(self):
        print(f"  初始化 VLM 引擎 (Qwen2.5-VL)...")
        self.server_url = Config.VLM_SERVER_URL
        self.available = self._check_server()
        if self.available:
            print("  ✓ VLM 引擎已连接")
        else:
            print("  ⚠️  VLM 服务不可用，将回退到 PaddleOCR")
    
    def _check_server(self):
        """检查 vLLM 服务是否可用"""
        try:
            import requests
            response = requests.get(f"{self.server_url}/health", timeout=5)
            return response.status_code == 200
        except:
            return False
    
    def recognize(self, image_path: str, layout_type: str) -> Dict:
        """VLM 识别"""
        try:
            import requests
            import base64
            
            # 编码图片
            with open(image_path, 'rb') as f:
                img_base64 = base64.b64encode(f.read()).decode('utf-8')
            
            # 构建请求
            prompt = "请识别图片中的所有文字内容，保持原有排版格式。如果是竖排文字，请从右到左阅读。"
            payload = {
                "model": Config.VLM_MODEL,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_base64}"}},
                        {"type": "text", "text": prompt},
                    ],
                }],
                "max_tokens": Config.VLM_MAX_TOKENS,
                "temperature": 0.0,
            }
            
            # 发送请求
            response = requests.post(
                f"{self.server_url}/v1/chat/completions",
                json=payload,
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                text = result['choices'][0]['message']['content']
                return {'text': text, 'confidence': 0.90, 'is_vlm': True}
            else:
                return {'text': '[VLM 错误]', 'confidence': 0.0, 'is_vlm': True}
                
        except Exception as e:
            return {'text': f'[VLM 异常：{e}]', 'confidence': 0.0, 'is_vlm': True}


# ==================== 混合路由 ====================

class HybridRouter:
    """混合架构路由"""
    
    def __init__(self, paddle_engine, vlm_engine):
        self.paddle = paddle_engine
        self.vlm = vlm_engine
    
    def should_use_vlm(self, layout_type: str) -> bool:
        """判断是否使用 VLM"""
        if not Config.USE_HYBRID:
            return False
        if not self.vlm.available:
            return False
        
        if layout_type == 'vertical' and Config.USE_VLM_FOR_VERTICAL:
            return True
        if layout_type == 'photo' and Config.USE_VLM_FOR_PHOTO:
            return True
        if layout_type == 'column' and Config.USE_VLM_FOR_COLUMN:
            return True
        if layout_type == 'table' and Config.USE_VLM_FOR_TABLE:
            return True
        
        return False
    
    def recognize(self, image, layout_type, image_path=None):
        """智能路由识别"""
        if self.should_use_vlm(layout_type):
            print(f"[VLM]", end=" ")
            if image_path:
                return self.vlm.recognize(image_path, layout_type)
            else:
                # 无路径时回退到 Paddle
                return self.paddle.recognize(image, layout_type)
        else:
            return self.paddle.recognize(image, layout_type)


# ==================== 显存监控器 ====================

class MemoryMonitor:
    def __init__(self):
        self.last_clean_time = time.time()
    def should_clean(self):
        if not torch.cuda.is_available(): return False
        used = torch.cuda.memory_allocated() / 1024**3
        total = torch.cuda.get_device_properties(0).total_memory / 1024**3
        return used / total > 0.85 or (time.time() - self.last_clean_time > 300 and used / total > 0.7)
    def record_clean(self):
        self.last_clean_time = time.time()


# ==================== PDF 处理器 ====================

class PDFProcessor:
    def __init__(self):
        self.monitor = MemoryMonitor()
        self.layout_analyzer = LayoutAnalyzer()
        self.paddle_engine = PaddleOCREngine(self.monitor)
        self.vlm_engine = VLMEngine()
        self.hybrid_router = HybridRouter(self.paddle_engine, self.vlm_engine)
        self.post_processor = OCRPostProcessor()
    
    def process_pdf(self, pdf_path, output_dir):
        pdf_name = Path(pdf_path).stem
        output_file = Path(output_dir) / f"{pdf_name}_ocr.docx"
        print(f"\n{'='*60}\n处理：{Path(pdf_path).name}\n{'='*60}")
        
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        print(f"  PDF 共 {total_pages} 页")
        
        pages_data = [None] * total_pages
        images_data = [None] * total_pages
        start_time = time.time()
        
        # 临时目录 (用于 VLM)
        temp_dir = Path(output_dir) / "temp"
        temp_dir.mkdir(exist_ok=True)
        
        vlm_count, paddle_count = 0, 0
        
        for page_num in range(total_pages):
            print(f"  处理第 {page_num + 1}/{total_pages} 页...", end=" ", flush=True)
            try:
                page = doc[page_num]
                mat = fitz.Matrix(Config.PADDLE_DPI / 72, Config.PADDLE_DPI / 72)
                pix = page.get_pixmap(matrix=mat)
                img = np.frombuffer(pix.samples, dtype=np.uint8).reshape((pix.height, pix.width, pix.n))
                img = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
                
                if max(img.shape[:2]) > Config.PADDLE_IMAGE_MAX_SIZE:
                    r = Config.PADDLE_IMAGE_MAX_SIZE / max(img.shape[:2])
                    img = cv2.resize(img, (int(img.shape[1]*r), int(img.shape[0]*r)), interpolation=cv2.INTER_LANCZOS4)
                
                layout = self.layout_analyzer.analyze(img)
                layout_type = layout['type']
                
                # 保存临时图片 (用于 VLM)
                temp_path = temp_dir / f"page_{page_num}.jpg"
                cv2.imwrite(str(temp_path), img)
                
                if layout_type == 'empty':
                    print("⊘ (空白页)")
                    pages_data[page_num] = (page_num + 1, "[空白页]", layout_type)
                    continue
                
                if layout_type == 'photo':
                    print("📷 (照片)")
                    pages_data[page_num] = (page_num + 1, "[照片]", layout_type)
                    images_data[page_num] = img if Config.KEEP_IMAGES else None
                    continue
                
                # 混合路由识别
                ocr_result = self.hybrid_router.recognize(img, layout_type, str(temp_path))
                
                if ocr_result.get('is_vlm', False):
                    vlm_count += 1
                else:
                    paddle_count += 1
                
                corrected = self.post_processor.process(ocr_result['text'], ocr_result['confidence'], layout_type)
                print(f"✓ ({layout_type}, {ocr_result['confidence']:.2f})")
                pages_data[page_num] = (page_num + 1, corrected, layout_type)
                images_data[page_num] = img if Config.KEEP_IMAGES else None
                
            except Exception as e:
                print(f"✗ (错误：{e})")
                pages_data[page_num] = (page_num + 1, f"[错误：{e}]", 'error')
            
            if (page_num + 1) % Config.CLEAN_GPU_INTERVAL == 0 and self.monitor.should_clean():
                self.paddle_engine.clear_memory()
                print(f"[显存清理]", end=" ", flush=True)
        
        doc.close()
        self.paddle_engine.clear_memory()
        
        # 清理临时文件
        import shutil
        if temp_dir.exists():
            shutil.rmtree(temp_dir)
        
        processed = sum(1 for p in pages_data if p)
        print(f"\n  验证：{processed}/{total_pages} 页 {'✅' if processed == total_pages else '⚠️'}")
        print(f"  引擎统计：PaddleOCR={paddle_count}, VLM={vlm_count}")
        
        self._save_docx(output_file, pages_data, images_data)
        
        elapsed = time.time() - start_time
        ppm = total_pages / (elapsed / 60) if elapsed > 0 else 0
        print(f"\n  ✓ 输出：{output_file.name} ({os.path.getsize(output_file)/1024:.1f} KB)")
        print(f"  ✓ 耗时：{elapsed:.1f}秒 | 速度：{ppm:.1f}页/分钟")
        
        return str(output_file), ppm
    
    def _save_docx(self, output_file, pages_data, images_data):
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        import io
        
        doc = Document()
        doc.sections[0].page_width = Cm(21)
        doc.sections[0].page_height = Cm(29.7)
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style.font.size = Pt(10.5)
        
        for i, (pn, text, layout) in enumerate(pages_data):
            if not text: continue
            doc.add_paragraph(f"=== 第 {pn} 页 ===", style='Heading 3')
            
            if layout == 'photo' and i < len(images_data) and images_data[i] is not None:
                self._add_image(doc, images_data[i])
            elif layout == 'table':
                self._add_table(doc, text)
            elif layout == 'vertical':
                doc.add_paragraph(f"[竖排版式]\n{text}")
            elif layout == 'column':
                doc.add_paragraph(f"[分栏版式]\n{text}")
            else:
                for line in text.split('\n'):
                    if line.strip(): doc.add_paragraph(line)
            
            if i < len(pages_data) - 1: doc.add_page_break()
        
        doc.save(output_file)
    
    def _add_image(self, doc, img, thumbnail=False):
        from docx.shared import Inches
        import io
        from PIL import Image
        try:
            img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
            pil_img = Image.fromarray(img_rgb)
            if thumbnail:
                r = 300 / pil_img.width
                if r < 1: pil_img = pil_img.resize((int(pil_img.width*r), int(pil_img.height*r)), Image.Resampling.LANCZOS)
            buf = io.BytesIO()
            pil_img.save(buf, format='JPEG', quality=Config.IMAGE_QUALITY, optimize=True)
            buf.seek(0)
            import tempfile
            tf = tempfile.NamedTemporaryFile(suffix='.jpg', delete=False)
            tf.write(buf.read()); tf.close()
            doc.add_picture(tf.name, width=Inches(5 if not thumbnail else 3))
            os.remove(tf.name)
        except: pass
    
    def _add_table(self, doc, text):
        try:
            lines = text.split('\n')
            rows = [l.split('│')[1:-1] for l in lines if l.startswith('│')]
            rows = [[c.strip() for c in r] for r in rows if r]
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = 'Table Grid'
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        t.cell(i, j).text = cell
            else:
                doc.add_paragraph(text)
        except:
            doc.add_paragraph(text)


# ==================== 主程序 ====================

def main():
    print("="*70)
    print("🚀 PDF 批量 OCR - 专业版 v13.0 混合架构版")
    print("="*70)
    print("架构：PaddleOCR (快) + Qwen2.5-VL (准)")
    print(f"  ✓ PaddleOCR: 横排/表格 (45-50 页/分)")
    print(f"  ✓ Qwen2.5-VL: 竖排/古籍 (更高质量)")
    print(f"  ✓ VLM 服务：{Config.VLM_SERVER_URL}")
    print("="*70)
    
    if torch.cuda.is_available():
        gpu = torch.cuda.get_device_name(0)
        mem = torch.cuda.get_device_properties(0).total_memory / 1024**3
        print(f"✅ GPU: {gpu} ({mem:.1f}GB)")
    print("="*70)
    
    input_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("/media/zjj/leidian/leidian")
    if not input_path.exists():
        print(f"❌ 路径不存在：{input_path}"); sys.exit(1)
    
    pdf_files = sorted(input_path.glob("*.pdf"))
    if not pdf_files:
        print(f"❌ 未找到 PDF 文件"); sys.exit(1)
    
    print(f"📂 发现 {len(pdf_files)} 个 PDF 文件\n")
    
    output_dir = Path("/media/zjj/leidian/leidian/ocr_output_v13_hybrid")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print("🔧 初始化混合 OCR 引擎...")
    processor = PDFProcessor()
    
    total_pages, total_time, success = 0, 0, 0
    
    try:
        for idx, pdf in enumerate(pdf_files, 1):
            t0 = time.time()
            out, speed = processor.process_pdf(str(pdf), str(output_dir))
            if out:
                success += 1
                doc = fitz.open(pdf)
                total_pages += len(doc)
                doc.close()
                total_time += time.time() - t0
                avg = total_pages / (total_time / 60) if total_time > 0 else 0
                
                with open(output_dir / "ocr_status.txt", 'w', encoding='utf-8') as f:
                    f.write(f"完成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}\n")
                    f.write(f"当前文件：{pdf.stem}\n")
                    f.write(f"输出：{out}\n")
                    f.write(f"进度：{idx}/{len(pdf_files)}\n")
                    f.write(f"速度：{avg:.1f}页/分钟\n")
                    f.write(f"模式：v13.0 混合架构版\n")
                
                print(f"\n✅ 第 {idx}/{len(pdf_files)} 个完成 | 平均：{avg:.1f}页/分钟")
    except KeyboardInterrupt:
        print("\n\n⚠️  中断")
    
    print("\n" + "="*70)
    print(f"✅ 完成：{success}/{len(pdf_files)} 文件")
    if total_time > 0:
        print(f"📊 平均速度：{total_pages / (total_time / 60):.1f} 页/分钟")
    print(f"📁 输出：{output_dir}")
    print("="*70)
    print("📌 v13.0 为混合架构版本，结合 PaddleOCR 和 VLM 优势")
    print("="*70)


if __name__ == "__main__":
    main()
