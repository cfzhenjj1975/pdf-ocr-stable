#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF 批量 OCR 处理脚本 - 专业版 v12.0 最终锁定版
配置：DPI=185 + 参数优化测试最优值
特性:
- DPI 185（平衡质量和速度）
- DET_THRESH=0.4（测试最优）
- BOX_THRESH=0.42（测试最优）
- UNCLIP_RATIO=1.3（测试最优）
- DROP_SCORE=0.55（测试最优）
- 严格表格检测
- 页码一致性保证
- 后处理校对

作者：OCR Team
日期：2026-02-25
版本：v12.0 FINAL Locked (DPI=185)
"""

import os, sys, gc, time, torch, warnings, cv2, numpy as np, fitz
from PIL import Image
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple
from ocr_postprocess import OCRPostProcessor

warnings.filterwarnings("ignore")

# ==================== v12.0 最终锁定配置 ====================

class Config:
    """OCR 配置 - v12.0 最终锁定 (DPI=185)"""
    
    GPU_MEMORY_GB = 13.0
    
    # DPI 固定 185（平衡质量和速度）
    DPI = 185
    IMAGE_MAX_SIZE = 1600
    
    # 参数优化测试最优值
    DET_DB_THRESH = 0.4        # 测试最优 (0.30-0.40 中最快)
    DET_DB_BOX_THRESH = 0.42   # 测试最优 (0.40-0.50 中最快)
    DET_DB_UNCLIP_RATIO = 1.3  # 测试最优 (1.0-1.4 中最快)
    DROP_SCORE = 0.55          # 测试最优 (满足质量要求)
    
    CLEAN_GPU_INTERVAL = 80
    MIN_CONFIDENCE = 0.5
    OUTPUT_FORMAT = "docx"
    KEEP_IMAGES = True
    KEEP_TEXT_IMAGES = False
    IMAGE_QUALITY = 90
    TABLE_WITH_BORDERS = True
    SKIP_EMPTY_PAGES = True
    EMPTY_PAGE_THRESHOLD = 0.02


# ==================== 版面分析 ====================

class LayoutAnalyzer:
    def analyze(self, image):
        h, w = image.shape[:2]
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        text_density = np.sum(binary > 0) / (h * w)
        if text_density < Config.EMPTY_PAGE_THRESHOLD:
            return {'type': 'empty', 'score': 0.95}
        if self._detect_photo(image, binary):
            return {'type': 'photo', 'score': 0.90}
        if self._detect_table_strict(binary):
            return {'type': 'table', 'score': 0.85}
        columns = self._detect_columns(binary)
        if len(columns) > 1:
            return {'type': 'column', 'columns': columns, 'score': 0.85}
        if self._detect_vertical_detailed(image, binary):
            return {'type': 'vertical', 'score': 0.80}
        return {'type': 'horizontal', 'score': 0.90}
    
    def _detect_photo(self, image, binary):
        h, w = image.shape[:2]
        edges = cv2.Canny(image, 50, 150)
        return np.sum(edges > 0) / (h * w) > 0.15 and np.sum(binary > 0) / (h * w) < 0.30
    
    def _detect_table_strict(self, binary):
        kernel_h = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
        kernel_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
        h_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel_h)
        v_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel_v)
        total = np.sum(binary > 0)
        if total == 0: return False
        h_ratio = np.sum(h_lines > 0) / total
        v_ratio = np.sum(v_lines > 0) / total
        if h_ratio < 0.05 or v_ratio < 0.05: return False
        return (h_ratio + v_ratio) > 0.20
    
    def _detect_columns(self, binary):
        h, w = binary.shape
        v_proj = np.sum(binary, axis=0)
        gaps, gap_start = [], None
        threshold = np.max(v_proj) * 0.1
        for i, val in enumerate(v_proj):
            if val < threshold:
                if gap_start is None: gap_start = i
            else:
                if gap_start is not None and i - gap_start > 50:
                    gaps.append((gap_start, i))
                gap_start = None
        if len(gaps) > 0:
            columns, last_end = [], 0
            for gs, ge in gaps:
                columns.append([last_end, 0, gs, h])
                last_end = ge
            columns.append([last_end, 0, w, h])
            return columns
        return []
    
    def _detect_vertical_detailed(self, image, binary):
        h, w = binary.shape
        h_proj = np.sum(binary, axis=1) / w
        v_proj = np.sum(binary, axis=0) / h
        if np.var(v_proj) < np.var(h_proj) * 1.3: return False
        num_labels, labels, stats, _ = cv2.connectedComponentsWithStats(binary, connectivity=8)
        if num_labels < 10: return False
        tall_count, total = 0, 0
        for i in range(1, num_labels):
            x, y, wb, hb, area = stats[i]
            if area < 50: continue
            if hb / wb > 1.2: tall_count += 1
            total += 1
        return np.var(v_proj) > np.var(h_proj) * 1.5 and tall_count / total > 0.5 if total > 0 else False


# ==================== OCR 引擎 ====================

class OCREngine:
    def __init__(self, monitor):
        from paddleocr import PaddleOCR
        self.monitor = monitor
        print(f"  初始化 OCR 引擎 (GPU:{Config.GPU_MEMORY_GB}GB, DPI:{Config.DPI})...")
        self.ocr_ch = self._create_engine('ch', True)
        self.ocr_cht = self._create_engine('chinese_cht', False)
        print("  ✓ OCR 引擎已加载")
    
    def _create_engine(self, lang, use_cls):
        from paddleocr import PaddleOCR
        return PaddleOCR(
            use_gpu=True, lang=lang, show_log=False,
            det=True, rec=True, cls=use_cls,
            gpu_mem=int(Config.GPU_MEMORY_GB * 1024),
            det_db_thresh=Config.DET_DB_THRESH,
            det_db_box_thresh=Config.DET_DB_BOX_THRESH,
            det_db_unclip_ratio=Config.DET_DB_UNCLIP_RATIO,
            drop_score=Config.DROP_SCORE,
            max_text_length=1000, use_space_char=True,
        )
    
    def recognize(self, image, layout_type):
        if layout_type == 'empty':
            return {'text': '[空白页]', 'confidence': 1.0}
        if layout_type == 'photo':
            return {'text': '[照片]', 'confidence': 1.0, 'is_photo': True}
        if layout_type == 'vertical':
            return self._recognize_vertical(image)
        return self._recognize_horizontal(image)
    
    def _recognize_horizontal(self, image):
        result = self.ocr_ch.ocr(image, cls=True)[0]
        if not result: return {'text': '', 'confidence': 0.0}
        texts, confs = [], []
        for item in result:
            if item:
                bbox, (text, conf) = item
                texts.append(text); confs.append(conf)
        return {'text': ' '.join(texts), 'confidence': sum(confs)/len(confs) if confs else 0.0}
    
    def _recognize_vertical(self, image):
        result = self.ocr_cht.ocr(image, cls=False)[0]
        if not result: return {'text': '', 'confidence': 0.0}
        words = []
        for item in result:
            if item:
                bbox, (text, conf) = item
                words.append({'text': text, 'bbox': bbox, 'conf': conf})
        words.sort(key=lambda x: (-x['bbox'][0][0], x['bbox'][0][1]))
        confs = [w['conf'] for w in words]
        return {'text': ''.join([w['text'] for w in words]), 'confidence': sum(confs)/len(confs) if confs else 0.0}
    
    def clear_memory(self):
        gc.collect()
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            torch.cuda.synchronize()
            self.monitor.record_clean()


# ==================== 显存监控器 ====================

class MemoryMonitor:
    def __init__(self):
        self.last_clean_time = time.time()
    def should_clean(self):
        if not torch.cuda.is_available(): return False
        used = torch.cuda.memory_allocated() / 1024**3
        total = torch.cuda.get_device_properties(0).total_memory / 1024**3
        return used / total > 0.85 or (time.time() - self.last_clean_time > 300 and used / total > 0.7)
    def record_clean(self):
        self.last_clean_time = time.time()


# ==================== PDF 处理器 ====================

class PDFProcessor:
    def __init__(self):
        self.monitor = MemoryMonitor()
        self.layout_analyzer = LayoutAnalyzer()
        self.ocr_engine = OCREngine(self.monitor)
        self.post_processor = OCRPostProcessor()
    
    def process_pdf(self, pdf_path, output_dir):
        pdf_name = Path(pdf_path).stem
        output_file = Path(output_dir) / f"{pdf_name}_ocr.docx"
        print(f"\n{'='*60}\n处理：{Path(pdf_path).name}\n{'='*60}")
        
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        print(f"  PDF 共 {total_pages} 页")
        
        pages_data = [None] * total_pages
        images_data = [None] * total_pages
        start_time = time.time()
        
        for page_num in range(total_pages):
            print(f"  处理第 {page_num + 1}/{total_pages} 页...", end=" ", flush=True)
            try:
                page = doc[page_num]
                mat = fitz.Matrix(Config.DPI / 72, Config.DPI / 72)
                pix = page.get_pixmap(matrix=mat)
                img = np.frombuffer(pix.samples, dtype=np.uint8).reshape((pix.height, pix.width, pix.n))
                img = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
                if max(img.shape[:2]) > Config.IMAGE_MAX_SIZE:
                    r = Config.IMAGE_MAX_SIZE / max(img.shape[:2])
                    img = cv2.resize(img, (int(img.shape[1]*r), int(img.shape[0]*r)), interpolation=cv2.INTER_LANCZOS4)
                
                layout = self.layout_analyzer.analyze(img)
                layout_type = layout['type']
                
                if layout_type == 'empty':
                    print("⊘ (空白页)")
                    pages_data[page_num] = (page_num + 1, "[空白页]", layout_type)
                    continue
                if layout_type == 'photo':
                    print("📷 (照片)")
                    pages_data[page_num] = (page_num + 1, "[照片]", layout_type)
                    images_data[page_num] = img if Config.KEEP_IMAGES else None
                    continue
                
                ocr_result = self.ocr_engine.recognize(img, layout_type)
                corrected = self.post_processor.process(ocr_result['text'], ocr_result['confidence'], layout_type)
                print(f"✓ ({layout_type}, {ocr_result['confidence']:.2f})")
                pages_data[page_num] = (page_num + 1, corrected, layout_type)
                images_data[page_num] = img if Config.KEEP_IMAGES else None
                
            except Exception as e:
                print(f"✗ (错误：{e})")
                pages_data[page_num] = (page_num + 1, f"[错误：{e}]", 'error')
            
            if (page_num + 1) % Config.CLEAN_GPU_INTERVAL == 0 and self.monitor.should_clean():
                self.ocr_engine.clear_memory()
                print(f"[显存清理]", end=" ", flush=True)
        
        doc.close()
        self.ocr_engine.clear_memory()
        
        processed = sum(1 for p in pages_data if p)
        print(f"\n  验证：{processed}/{total_pages} 页 {'✅' if processed == total_pages else '⚠️'}")
        
        self._save_docx(output_file, pages_data, images_data)
        
        elapsed = time.time() - start_time
        ppm = total_pages / (elapsed / 60) if elapsed > 0 else 0
        print(f"\n  ✓ 输出：{output_file.name} ({os.path.getsize(output_file)/1024:.1f} KB)")
        print(f"  ✓ 耗时：{elapsed:.1f}秒 | 速度：{ppm:.1f}页/分钟")
        
        return str(output_file), ppm
    
    def _save_docx(self, output_file, pages_data, images_data):
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        import io
        
        doc = Document()
        doc.sections[0].page_width = Cm(21)
        doc.sections[0].page_height = Cm(29.7)
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style.font.size = Pt(10.5)
        
        for i, (pn, text, layout) in enumerate(pages_data):
            if not text: continue
            doc.add_paragraph(f"=== 第 {pn} 页 ===", style='Heading 3')
            
            if layout == 'photo' and i < len(images_data) and images_data[i] is not None:
                self._add_image(doc, images_data[i])
            elif layout == 'table':
                self._add_table(doc, text)
            elif layout == 'vertical':
                doc.add_paragraph(f"[竖排版式]\n{text}")
            elif layout == 'column':
                doc.add_paragraph(f"[分栏版式]\n{text}")
            else:
                for line in text.split('\n'):
                    if line.strip(): doc.add_paragraph(line)
            
            if i < len(pages_data) - 1: doc.add_page_break()
        
        doc.save(output_file)
    
    def _add_image(self, doc, img, thumbnail=False):
        from docx.shared import Inches
        import io
        from PIL import Image
        try:
            img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
            pil_img = Image.fromarray(img_rgb)
            if thumbnail:
                r = 300 / pil_img.width
                if r < 1: pil_img = pil_img.resize((int(pil_img.width*r), int(pil_img.height*r)), Image.Resampling.LANCZOS)
            buf = io.BytesIO()
            pil_img.save(buf, format='JPEG', quality=Config.IMAGE_QUALITY, optimize=True)
            buf.seek(0)
            import tempfile
            tf = tempfile.NamedTemporaryFile(suffix='.jpg', delete=False)
            tf.write(buf.read()); tf.close()
            doc.add_picture(tf.name, width=Inches(5 if not thumbnail else 3))
            os.remove(tf.name)
        except: pass
    
    def _add_table(self, doc, text):
        try:
            lines = text.split('\n')
            rows = [l.split('│')[1:-1] for l in lines if l.startswith('│')]
            rows = [[c.strip() for c in r] for r in rows if r]
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = 'Table Grid'
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        t.cell(i, j).text = cell
            else:
                doc.add_paragraph(text)
        except:
            doc.add_paragraph(text)


# ==================== 主程序 ====================

def main():
    print("="*70)
    print("🚀 PDF 批量 OCR - 专业版 v12.0 最终锁定版 (DPI=185)")
    print("="*70)
    print("配置：DPI=185 + 参数优化测试最优值")
    print(f"  ✓ DPI: {Config.DPI}")
    print(f"  ✓ DET_THRESH: {Config.DET_DB_THRESH}")
    print(f"  ✓ BOX_THRESH: {Config.DET_DB_BOX_THRESH}")
    print(f"  ✓ UNCLIP_RATIO: {Config.DET_DB_UNCLIP_RATIO}")
    print(f"  ✓ DROP_SCORE: {Config.DROP_SCORE}")
    print("="*70)
    
    if torch.cuda.is_available():
        gpu = torch.cuda.get_device_name(0)
        mem = torch.cuda.get_device_properties(0).total_memory / 1024**3
        print(f"✅ GPU: {gpu} ({mem:.1f}GB)")
    print("="*70)
    
    input_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("/media/zjj/leidian/leidian")
    if not input_path.exists():
        print(f"❌ 路径不存在：{input_path}"); sys.exit(1)
    
    pdf_files = sorted(input_path.glob("*.pdf"))
    if not pdf_files:
        print(f"❌ 未找到 PDF 文件"); sys.exit(1)
    
    print(f"📂 发现 {len(pdf_files)} 个 PDF 文件\n")
    
    output_dir = Path("/media/zjj/leidian/leidian/ocr_output_v12_final")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print("🔧 初始化 OCR 引擎...")
    processor = PDFProcessor()
    
    total_pages, total_time, success = 0, 0, 0
    
    try:
        for idx, pdf in enumerate(pdf_files, 1):
            t0 = time.time()
            out, speed = processor.process_pdf(str(pdf), str(output_dir))
            if out:
                success += 1
                doc = fitz.open(pdf)
                total_pages += len(doc)
                doc.close()
                total_time += time.time() - t0
                avg = total_pages / (total_time / 60) if total_time > 0 else 0
                
                with open(output_dir / "ocr_status.txt", 'w', encoding='utf-8') as f:
                    f.write(f"完成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}\n")
                    f.write(f"当前文件：{pdf.stem}\n")
                    f.write(f"输出：{out}\n")
                    f.write(f"进度：{idx}/{len(pdf_files)}\n")
                    f.write(f"速度：{avg:.1f}页/分钟\n")
                    f.write(f"模式：v12.0 最终锁定版 (DPI={Config.DPI})\n")
                
                print(f"\n✅ 第 {idx}/{len(pdf_files)} 个完成 | 平均：{avg:.1f}页/分钟")
    except KeyboardInterrupt:
        print("\n\n⚠️  中断")
    
    print("\n" + "="*70)
    print(f"✅ 完成：{success}/{len(pdf_files)} 文件")
    if total_time > 0:
        print(f"📊 平均速度：{total_pages / (total_time / 60):.1f} 页/分钟")
    print(f"📁 输出：{output_dir}")
    print("="*70)
    print("📌 v12.0 为最终锁定版本，参数已通过测试优化确定")
    print("="*70)


if __name__ == "__main__":
    main()
